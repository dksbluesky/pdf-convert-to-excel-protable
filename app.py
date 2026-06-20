import os
import re
import json
import datetime
import tempfile
import concurrent.futures
from urllib.parse import quote
from io import BytesIO
from flask import Flask, render_template, request, jsonify, send_file
import pdfplumber
import pandas as pd
import requests
import openpyxl

try:
    from pdf2image import convert_from_bytes
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    pass

try:
    from pdf2docx import Converter as PDF2Docx
    PDF2DOCX_AVAILABLE = True
except ImportError:
    PDF2DOCX_AVAILABLE = False

GEMINI_API_KEY = os.environ.get('GEMINI_API_KEY', '')
GEMINI_AVAILABLE = bool(GEMINI_API_KEY)
if GEMINI_AVAILABLE:
    try:
        import google.generativeai as genai
        genai.configure(api_key=GEMINI_API_KEY)
    except ImportError:
        GEMINI_AVAILABLE = False

GEMINI_MAX_BYTES = 20 * 1024 * 1024  # 20 MB inline limit

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 300 * 1024 * 1024  # 300 MB total (AI mode batch uploads)


# ──────────────────────────────────────────────
# Analysis
# ──────────────────────────────────────────────
def analyze_pdf(pdf_bytes: bytes) -> dict:
    table_pages = text_pages = total_pages = 0
    try:
        with pdfplumber.open(BytesIO(pdf_bytes)) as pdf:
            total_pages = len(pdf.pages)
            for page in pdf.pages[:min(5, total_pages)]:
                tables = page.extract_tables() or page.extract_tables(
                    table_settings={"vertical_strategy": "text", "horizontal_strategy": "text"})
                if tables:
                    table_pages += 1
                if (page.extract_text() or "").strip():
                    text_pages += 1
    except Exception:
        pass

    if text_pages == 0 and table_pages == 0:
        return {"format": "word", "badge": "🔍 掃描版 PDF",
                "reason": "未偵測到可擷取的文字，建議啟用 OCR 再轉為 Word。", "use_ocr": True}
    elif table_pages >= text_pages * 0.6:
        return {"format": "excel", "badge": "📊 表格型 PDF",
                "reason": f"前 {min(5, total_pages)} 頁中有 {table_pages} 頁含表格，最適合匯出為 Excel。", "use_ocr": False}
    elif table_pages > 0:
        return {"format": "word", "badge": "📝 混合型 PDF",
                "reason": f"文字與表格混合（{table_pages} 頁有表格），建議轉為 Word。", "use_ocr": False}
    else:
        return {"format": "markdown", "badge": "📄 純文字 PDF",
                "reason": "純文字內容，Markdown 最輕量且易讀。", "use_ocr": False}


# ──────────────────────────────────────────────
# Extraction
# ──────────────────────────────────────────────
def _unique_cols(headers: list) -> list:
    seen: dict = {}
    out = []
    for h in headers:
        key = h if h else "col"
        if key in seen:
            seen[key] += 1
            out.append(f"{key}_{seen[key]}")
        else:
            seen[key] = 0
            out.append(key)
    return out


def extract_tables(pdf_bytes: bytes) -> dict:
    result = {}
    with pdfplumber.open(BytesIO(pdf_bytes)) as pdf:
        for i, page in enumerate(pdf.pages):
            tables = page.extract_tables(table_settings={
                "vertical_strategy": "lines", "horizontal_strategy": "lines"})
            if not tables:
                tables = page.extract_tables(table_settings={
                    "vertical_strategy": "text", "horizontal_strategy": "text"})
            if tables:
                dfs = []
                for t in tables:
                    cleaned = [[c if c is not None else "" for c in row] for row in t]
                    if not cleaned:
                        continue
                    if len(cleaned) > 1:
                        df = pd.DataFrame(cleaned[1:], columns=_unique_cols(cleaned[0]))
                    else:
                        df = pd.DataFrame(cleaned)
                    dfs.append(df.astype(str).replace("None", ""))
                if dfs:
                    try:
                        result[f"Page {i+1}"] = pd.concat(dfs, ignore_index=True)
                    except Exception:
                        normed = [df.rename(columns=str) for df in dfs]
                        result[f"Page {i+1}"] = pd.concat(normed, ignore_index=True)
    return result


def extract_ocr(pdf_bytes: bytes) -> dict:
    images = convert_from_bytes(pdf_bytes, dpi=300)
    result = {}
    for i, img in enumerate(images):
        try:
            text = pytesseract.image_to_string(img, lang="chi_tra+eng")
        except Exception:
            text = pytesseract.image_to_string(img)
        if text.strip():
            result[f"Page {i+1}"] = text.strip()
    return result


def extract_text_pages(pdf_bytes: bytes) -> dict:
    result = {}
    with pdfplumber.open(BytesIO(pdf_bytes)) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            if text.strip():
                result[f"Page {i+1}"] = text.strip()
    return result


# ──────────────────────────────────────────────
# Build output
# ──────────────────────────────────────────────
def build_excel(data: dict, notes: dict = None) -> bytes:
    out = BytesIO()
    with pd.ExcelWriter(out, engine="openpyxl") as writer:
        for sheet, df in data.items():
            sheet_name = sheet[:31]
            note = (notes or {}).get(sheet)
            startrow = 1 if note else 0
            df_out, numeric_cols = _numericize_for_excel(df)
            df_out.to_excel(writer, sheet_name=sheet_name, index=False, startrow=startrow)
            if note:
                writer.sheets[sheet_name]["A1"] = note
            if numeric_cols:
                ws = writer.sheets[sheet_name]
                header_row = startrow + 1
                cols = list(df_out.columns)
                for col_name in numeric_cols:
                    col_idx = cols.index(col_name) + 1
                    for r in range(header_row + 1, header_row + 1 + len(df_out)):
                        cell = ws.cell(row=r, column=col_idx)
                        if cell.value not in (None, ""):
                            cell.number_format = "#,##0"
    return out.getvalue()


_RATE_NOTE_PREFIX = "匯率："


def read_excel_sheets(file_bytes: bytes) -> dict:
    """Read all sheets from an uploaded Excel file, the way build_excel()'s own
    output needs to be read back during "加入現有 Excel": when a sheet has a
    currency rate-note in cell A1 (write side: build_excel's startrow=1 shift),
    the real header row is row 2, not row 1 — reading with the default header=0
    would treat the note text as a column header and shift every real column
    name to "Unnamed: N", silently corrupting that sheet's data on merge."""
    wb = openpyxl.load_workbook(BytesIO(file_bytes), read_only=True, data_only=True)
    sheets = {}
    try:
        for name in wb.sheetnames:
            ws = wb[name]
            a1 = ws.cell(row=1, column=1).value
            b1 = ws.cell(row=1, column=2).value
            has_note = isinstance(a1, str) and a1.startswith(_RATE_NOTE_PREFIX) and b1 in (None, "")
            sheets[name] = pd.read_excel(BytesIO(file_bytes), sheet_name=name, header=1 if has_note else 0)
    finally:
        wb.close()
    return sheets


def build_word(pdf_bytes: bytes, is_ocr: bool, ocr_data: dict = None) -> bytes:
    if is_ocr and ocr_data:
        doc = Document()
        doc.add_heading("PDF 轉換結果（OCR）", level=0)
        for page, text in ocr_data.items():
            doc.add_heading(page, level=1)
            doc.add_paragraph(text)
        out = BytesIO()
        doc.save(out)
        return out.getvalue()

    if PDF2DOCX_AVAILABLE:
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            f.write(pdf_bytes)
            tmp_pdf = f.name
        tmp_docx = tmp_pdf.replace(".pdf", ".docx")
        try:
            cv = PDF2Docx(tmp_pdf)
            cv.convert(tmp_docx)
            cv.close()
            with open(tmp_docx, "rb") as f:
                return f.read()
        finally:
            os.unlink(tmp_pdf)
            if os.path.exists(tmp_docx):
                os.unlink(tmp_docx)

    doc = Document()
    doc.add_heading("PDF 轉換結果", level=0)
    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def build_word_editable(pdf_bytes: bytes, table_data: dict = None) -> bytes:
    pages = extract_text_pages(pdf_bytes)
    doc = Document()
    if pages:
        for label, text in pages.items():
            doc.add_heading(label, level=1)
            for line in text.splitlines():
                if line.strip():
                    doc.add_paragraph(line.strip())
            doc.add_paragraph()
    elif table_data:
        for label, df in table_data.items():
            doc.add_heading(label, level=1)
            for _, row in df.iterrows():
                line = "　".join(str(v) for v in row if str(v).strip())
                if line.strip():
                    doc.add_paragraph(line)
            doc.add_paragraph()
    else:
        doc.add_paragraph("⚠️ 此 PDF 無法擷取文字，請改用 OCR 模式。")
    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def _append_docx(existing_bytes: bytes, new_bytes: bytes) -> bytes:
    base_doc = Document(BytesIO(existing_bytes))
    new_doc = Document(BytesIO(new_bytes))
    base_doc.add_page_break()
    for element in list(new_doc.element.body):
        if element.tag == qn("w:sectPr"):
            continue
        base_doc.element.body.append(element)
    out = BytesIO()
    base_doc.save(out)
    return out.getvalue()


def build_markdown(data: dict, is_ocr: bool) -> str:
    parts = ["# PDF 轉換結果\n"]
    for page, content in data.items():
        parts.append(f"\n## {page}\n")
        parts.append(content if is_ocr else content.to_markdown(index=False))
        parts.append("\n")
    return "\n".join(parts)


def build_json(data: dict, is_ocr: bool) -> str:
    payload = data if is_ocr else {p: df.to_dict(orient="records") for p, df in data.items()}
    return json.dumps(payload, ensure_ascii=False, indent=2)


# ──────────────────────────────────────────────
# Routes — general mode
# ──────────────────────────────────────────────
@app.route("/")
def index():
    return render_template("index.html",
                           ocr_available=OCR_AVAILABLE,
                           gemini_available=GEMINI_AVAILABLE)


@app.route("/analyze", methods=["POST"])
def analyze_route():
    if "pdf" not in request.files:
        return jsonify({"error": "No file"}), 400
    rec = analyze_pdf(request.files["pdf"].read())
    return jsonify(rec)


@app.route("/convert", methods=["POST"])
def convert_route():
    if "pdf" not in request.files:
        return jsonify({"error": "No file"}), 400

    pdf_bytes = request.files["pdf"].read()
    fmt = request.form.get("format", "excel")
    use_ocr = request.form.get("ocr") == "true"
    word_editable = request.form.get("word_editable") == "true"
    base_name = request.form.get("filename", "converted")
    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")

    existing_file = request.files.get("existing_file")
    existing_bytes = existing_file.read() if existing_file and existing_file.filename else None
    existing_label = existing_file.filename.rsplit(".", 1)[0] if existing_bytes else None

    if use_ocr and OCR_AVAILABLE:
        data = extract_ocr(pdf_bytes)
        is_ocr = True
    else:
        data = extract_tables(pdf_bytes)
        is_ocr = False

    if fmt == "excel":
        if not data:
            return jsonify({"error": "no_data"}), 422
        if existing_bytes:
            try:
                existing_sheets = read_excel_sheets(existing_bytes)
            except Exception:
                existing_sheets = {}
            used = set(existing_sheets.keys())
            all_sheets = dict(existing_sheets)
            multi_page = len(data) > 1
            for label, df in data.items():
                raw_name = f"{base_name}_{label}" if multi_page else base_name
                all_sheets[_excel_safe_sheet_name(raw_name, used)] = df
            file_bytes = build_excel(all_sheets)
            filename = f"{existing_label}_更新_{ts}.xlsx"
        else:
            file_bytes = build_excel(data)
            filename = f"{base_name}_{ts}.xlsx"
        mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

    elif fmt == "word":
        if word_editable:
            new_bytes = build_word_editable(pdf_bytes, table_data=data if not is_ocr else None)
        else:
            new_bytes = build_word(pdf_bytes, is_ocr, ocr_data=data if is_ocr else None)
        if existing_bytes:
            try:
                file_bytes = _append_docx(existing_bytes, new_bytes)
                filename = f"{existing_label}_更新_{ts}.docx"
            except Exception:
                file_bytes = new_bytes
                filename = f"{base_name}_{ts}.docx"
        else:
            file_bytes = new_bytes
            filename = f"{base_name}_{ts}.docx"
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    elif fmt == "markdown":
        if not data:
            return jsonify({"error": "no_data"}), 422
        new_md = build_markdown(data, is_ocr)
        if existing_bytes:
            try:
                old_md = existing_bytes.decode("utf-8")
            except Exception:
                old_md = ""
            file_bytes = (old_md.rstrip() + "\n\n---\n\n" + new_md).encode("utf-8")
            filename = f"{existing_label}_更新_{ts}.md"
        else:
            file_bytes = new_md.encode("utf-8")
            filename = f"{base_name}_{ts}.md"
        mime = "text/markdown"

    elif fmt == "json":
        if not data:
            return jsonify({"error": "no_data"}), 422
        new_payload = data if is_ocr else {p: df.to_dict(orient="records") for p, df in data.items()}
        if existing_bytes:
            try:
                old_payload = json.loads(existing_bytes.decode("utf-8"))
                if not isinstance(old_payload, dict):
                    old_payload = {"原始資料": old_payload}
            except Exception:
                old_payload = {}
            used = set(old_payload.keys())
            merged = dict(old_payload)
            for key, value in new_payload.items():
                safe_key, i = key, 1
                while safe_key in used:
                    safe_key = f"{key}_{i}"
                    i += 1
                used.add(safe_key)
                merged[safe_key] = value
            file_bytes = json.dumps(merged, ensure_ascii=False, indent=2).encode("utf-8")
            filename = f"{existing_label}_更新_{ts}.json"
        else:
            file_bytes = json.dumps(new_payload, ensure_ascii=False, indent=2).encode("utf-8")
            filename = f"{base_name}_{ts}.json"
        mime = "application/json"

    else:
        return jsonify({"error": "Unknown format"}), 400

    return send_file(BytesIO(file_bytes), mimetype=mime,
                     as_attachment=True, download_name=filename)


# ──────────────────────────────────────────────
# Routes — AI mode (Gemini)
# ──────────────────────────────────────────────
PROMPT_AI_PLAIN = """你是一個專業的資料輸入員。請將這份圖片或 PDF 中的表格轉換為純文字資料。

【嚴格規則】
1. 在輸出表格之前，請先單獨輸出一行中繼資料，格式為：META###日期###店家或場所名稱。日期請盡量判斷並轉換為 YYYY-MM-DD 格式（只取日期，不需要時間），完全無法判斷時請留空；店家/場所名稱請填入收據上的店名或場所全名。這一行只會出現一次，且只包含這三個部分，輸出完這一行之後才開始輸出表格本身。
2. 中繼資料那一行之後，每一欄之間請使用 "###" 作為分隔符號。即使表格只有兩、三欄，也一定要在每一欄之間加上 "###"，絕對不可以把一整列當作一段純文字直接輸出而不分欄。
3. 每一列資料換一行。
4. 表格的第一行必須是表頭，欄位數請以「真正的商品/項目表格」的欄數為準。
5. 表格上方若有店名、地址、營業資訊、桌號、訂單時間、收據編號等與商品項目無關的資訊，這些已經在第 1 點的中繼資料處理過，不要再把它們當成表格的列輸出。只擷取「有多筆品名/數量/金額之類欄位」的那個表格本身。
6. 若某一列在某欄沒有對應的內容（例如備註列在數量、單價欄沒有資料），該儲存格請留空，不要填入任何文字、符號或占位字元（例如「無」、「-」、「叢」等）。
7. 不要輸出任何 Markdown 標記，只要純文字。
8. 金額請保留千分位符號，不要隨意移除。
9. 若遇到跨頁，請自動合併。
10. 底部若有付款條件、稅金、總計等資訊，請整理在表格最下方的列，並把該資訊文字放在「品名／項目」欄；但「金額」欄請留空，不要填入小計、總計、含稅金額等任何數字，因為這些列是文字說明列，不是逐項商品，重複填入金額會導致與表格上方各項金額加總後重複計算。"""

PROMPT_AI_TRANSLATE = """你是一個專業的雙語資料輸入員。請將這份圖片或 PDF 中的表格轉換為純文字資料，並在同一儲存格內附上中文翻譯。

【嚴格規則】
1. 在輸出表格之前，請先單獨輸出一行中繼資料，格式為：META###日期###店家或場所名稱。日期請盡量判斷並轉換為 YYYY-MM-DD 格式（只取日期，不需要時間），完全無法判斷時請留空；店家/場所名稱若為非中文，請依規則 5 的格式加上中文翻譯（原文 / 中文）。這一行只會出現一次，且只包含這三個部分，輸出完這一行之後才開始輸出表格本身。
2. 中繼資料那一行之後，每一欄之間請使用 "###" 作為分隔符號。即使表格只有兩、三欄，也一定要在每一欄之間加上 "###"，絕對不可以把一整列當作一段純文字直接輸出而不分欄。
3. 每一列資料換一行，表格的第一行必須是表頭，欄位數請以「真正的商品/項目表格」的欄數為準。
4. 表格上方若有店名、地址、營業資訊、桌號、訂單時間、收據編號等與商品項目無關的資訊，這些已經在第 1 點的中繼資料處理過，不要再把它們當成表格的列輸出。只擷取「有多筆品名/數量/金額之類欄位」的那個表格本身。
5. 請以「儲存格」為單位逐一判斷（包含表頭儲存格與第 1 點的店家/場所名稱），不要只看欄位標題或其他列的內容：只要某個儲存格包含非中文的文字說明（例如店名、地址、品名、付款方式、備註、表頭名稱等），請在該儲存格加上中文翻譯，翻譯一律使用「繁體中文」，絕對不可以翻譯成英文或其他語言。格式必須嚴格遵守：「原文」+ 一個半角空格 + 一個半角斜線「/」+ 一個半角空格 + 「繁體中文翻譯」。斜線前後都一定要有空格，不可省略，也不可使用全角斜線「／」。正確範例：「아메리카노 / 美式咖啡」、「품목 / 品名」。錯誤範例（缺空格，禁止）：「아메리카노/美式咖啡」。錯誤範例（翻成英文，禁止）：「아메리카노 / Americano」。
6. 特別注意：店名、分店名稱、付款方式（如信用卡、現金）等即使出現在看起來像「數值欄」的位置，仍然是文字內容，必須加上翻譯，不可因為欄位標題像數字欄而跳過。
7. 純數字、金額、日期、編號儲存格維持原樣即可，不需要翻譯，也不要加上斜線。
8. 不要新增任何欄位，輸出的欄位數量必須與原表格相同。
9. 若某一列在某欄沒有對應的內容（例如備註列在數量、單價欄沒有資料），該儲存格請留空，不要填入任何文字、符號或占位字元（例如「無」、「-」、「叢」等）。
10. 不要輸出任何 Markdown 標記，只要純文字。
11. 金額請保留千分位符號，不要隨意移除。
12. 若遇到跨頁，請自動合併。
13. 底部若有付款條件、稅金、總計等資訊，請整理在表格最下方的列，並把該資訊文字放在「品名／項目」欄，文字部分同樣依規則 5、6、7 處理；但「金額」欄請留空，不要填入小計、總計、含稅金額等任何數字，因為這些列是文字說明列，不是逐項商品，重複填入金額會導致與表格上方各項金額加總後重複計算。"""


def _gemini_mime(filename: str):
    fname = filename.lower()
    if fname.endswith(".pdf"):
        return "application/pdf"
    elif fname.endswith((".jpg", ".jpeg")):
        return "image/jpeg"
    elif fname.endswith(".png"):
        return "image/png"
    return None


def _excel_safe_sheet_name(name: str, used: set) -> str:
    invalid = "[]:*?/\\"
    clean = "".join(c for c in name if c not in invalid).strip() or "Sheet"
    clean = clean[:31]
    base, i = clean, 1
    while clean in used:
        suffix = f"_{i}"
        clean = base[:31 - len(suffix)] + suffix
        i += 1
    used.add(clean)
    return clean


NTD_COL = "NT$ / 金額"
AMOUNT_KEYWORDS = ["金額", "金额", "금액", "amount", "total", "合計", "总额", "総額", "subtotal"]


def _is_amount_col(name) -> bool:
    if name == NTD_COL:
        return False
    low = str(name).lower()
    return any(kw.lower() in low for kw in AMOUNT_KEYWORDS)


def _find_amount_col(columns):
    for c in columns:
        if _is_amount_col(c):
            return c
    return None


def _coalesce_amount_columns(df: pd.DataFrame) -> pd.DataFrame:
    """Different invoices may translate their amount header slightly differently
    (e.g. "금액 / 金額" vs "총금액 / 總金額"), which pd.concat treats as separate
    columns. Unify them into one so totals sum correctly."""
    amount_cols = [c for c in df.columns if _is_amount_col(c)]
    if len(amount_cols) <= 1:
        return df
    out = df.copy()
    primary = amount_cols[0]
    # One side may be a real float64 column (read back from a previously
    # written Excel file, after the numeric-formatting fix) while the other
    # is still a string column (freshly extracted, e.g. "16,000") — assigning
    # strings into a strict float64 column raises in pandas 2.x. Widen to
    # object first; _to_numeric() parses either representation later anyway.
    out[primary] = out[primary].astype(object)
    for c in amount_cols[1:]:
        is_blank = out[primary].isna() | (out[primary].astype(str).str.strip() == "")
        out.loc[is_blank, primary] = out.loc[is_blank, c]
        out = out.drop(columns=[c])
    return out


def _to_numeric(series: pd.Series) -> pd.Series:
    def parse(v):
        if pd.isna(v):
            return None
        if isinstance(v, (int, float)):
            return float(v)
        m = re.search(r"-?\d[\d,]*\.?\d*", str(v))
        return float(m.group().replace(",", "")) if m else None
    return series.map(parse)


NUMERIC_HEADER_KEYWORDS = AMOUNT_KEYWORDS + [
    "數量", "数量", "수량", "quantity", "qty",
    "單價", "单价", "단가", "unit price", "price",
]
_NUMERIC_CELL_RE = re.compile(r"^-?[\d,]+\.?\d*$")


def _is_numeric_header(name) -> bool:
    low = str(name).lower()
    return any(kw.lower() in low for kw in NUMERIC_HEADER_KEYWORDS)


def _numericize_for_excel(df: pd.DataFrame):
    """Quantity/price/amount columns are kept as text strings (e.g. "9,000")
    so the original thousands-separator formatting is preserved verbatim, but
    that makes Excel treat them as text — no native AutoSum/status-bar totals.
    Convert columns that are unambiguously numeric (by header keyword AND every
    non-blank cell matching a plain number pattern) into real numbers with an
    Excel thousands-separator number format, so they look the same but sum natively."""
    out = df.copy()
    numeric_cols = []
    for c in out.columns:
        if c == "來源檔案" or not _is_numeric_header(c):
            continue
        vals = out[c].astype(str).str.strip()
        non_blank = vals[(vals != "") & (vals.str.lower() != "nan")]
        if non_blank.empty or not non_blank.str.match(_NUMERIC_CELL_RE).all():
            continue
        out[c] = _to_numeric(out[c])
        numeric_cols.append(c)
    return out, numeric_cols


def _add_ntd_column(df: pd.DataFrame, rate: float) -> pd.DataFrame:
    amount_col = _find_amount_col(df.columns)
    if not amount_col or NTD_COL in df.columns:
        return df
    numeric = _to_numeric(df[amount_col])
    ntd_values = numeric.map(lambda v: round(v * rate) if pd.notna(v) else "")
    out = df.copy()
    out.insert(list(df.columns).index(amount_col) + 1, NTD_COL, ntd_values)
    return out


def _per_file_subtotal_row(df: pd.DataFrame) -> dict:
    row = {c: "" for c in df.columns}
    row[df.columns[0]] = "小計"
    amount_col = _find_amount_col(df.columns)
    if amount_col:
        row[amount_col] = _to_numeric(df[amount_col]).sum()
    if NTD_COL in df.columns:
        row[NTD_COL] = _to_numeric(df[NTD_COL]).sum()
    return row


# ──────────────────────────────────────────────
# Raw Sheet — normalized one-row-per-item view across all invoices, built for
# pivot tables / dashboards (fixed columns regardless of each receipt's own
# layout/language), grouped by date with per-date and grand-total rows.
# ──────────────────────────────────────────────
RAW_SHEET_NAME = "Raw Sheet"
RAW_DATE_COL = "日期"
RAW_STORE_COL = "店家/場所"
RAW_ITEM_COL = "品名"
RAW_DESC_COL = "說明"
RAW_QTY_COL = "數量"
RAW_PRICE_COL = "單價"
RAW_AMOUNT_COL = "金額"
RAW_COLUMNS = [RAW_DATE_COL, RAW_STORE_COL, RAW_ITEM_COL, RAW_DESC_COL,
               RAW_QTY_COL, RAW_PRICE_COL, RAW_AMOUNT_COL]

QTY_KEYWORDS = ["數量", "数量", "수량", "quantity", "qty"]
PRICE_KEYWORDS = ["單價", "单价", "단가", "unit price", "price"]


def _find_col(columns, keywords, exclude=None):
    for c in columns:
        if exclude and c in exclude:
            continue
        low = str(c).lower()
        if any(kw.lower() in low for kw in keywords):
            return c
    return None


def _normalize_invoice_rows(df: pd.DataFrame, date_str: str, store_str: str) -> pd.DataFrame:
    """Maps one invoice's own (varying) table into the fixed Raw Sheet schema.
    Rows with neither a quantity nor an amount are footer/note text (tax
    notice, payment method, thank-you line) rather than purchased items, and
    are skipped — they're not part of what was actually bought."""
    item_col = df.columns[0]
    amount_col = _find_amount_col(df.columns)
    qty_col = _find_col(df.columns, QTY_KEYWORDS)
    price_col = _find_col(df.columns, PRICE_KEYWORDS, exclude={amount_col})
    has_ntd = NTD_COL in df.columns

    out_rows = []
    for _, r in df.iterrows():
        qty_val = r[qty_col] if qty_col else ""
        amount_val = r[amount_col] if amount_col else ""
        qty_blank = qty_val is None or (isinstance(qty_val, float) and pd.isna(qty_val)) or str(qty_val).strip() == ""
        amt_blank = amount_val is None or (isinstance(amount_val, float) and pd.isna(amount_val)) or str(amount_val).strip() == ""
        if qty_blank and amt_blank:
            continue
        row = {
            RAW_DATE_COL: date_str,
            RAW_STORE_COL: store_str,
            RAW_ITEM_COL: str(r[item_col]).strip(),
            RAW_DESC_COL: "",
            RAW_QTY_COL: qty_val,
            RAW_PRICE_COL: r[price_col] if price_col else "",
            RAW_AMOUNT_COL: amount_val,
        }
        if has_ntd:
            row[NTD_COL] = r[NTD_COL]
        out_rows.append(row)
    cols = RAW_COLUMNS + ([NTD_COL] if has_ntd else [])
    return pd.DataFrame(out_rows, columns=cols)


def _strip_raw_marker_rows(df: pd.DataFrame) -> pd.DataFrame:
    if df.empty or RAW_DATE_COL not in df.columns:
        return df
    mask = ~df[RAW_DATE_COL].astype(str).str.startswith(("小計", "總計"))
    return df[mask].reset_index(drop=True)


def _raw_total_row(df: pd.DataFrame, label: str) -> dict:
    clean = _strip_raw_marker_rows(df)
    row = {c: "" for c in df.columns}
    row[RAW_DATE_COL] = label
    if RAW_AMOUNT_COL in df.columns:
        row[RAW_AMOUNT_COL] = _to_numeric(clean[RAW_AMOUNT_COL]).sum()
    if NTD_COL in df.columns:
        row[NTD_COL] = _to_numeric(clean[NTD_COL]).sum()
    return row


def _build_raw_sheet(new_rows: pd.DataFrame, existing_raw: pd.DataFrame = None) -> pd.DataFrame:
    """Merges this round's normalized rows into the existing Raw Sheet (if any),
    strips stale 小計/總計 rows from the old data, then regroups everything by
    date so two receipts sharing a date end up under one subtotal — this is
    what lets "加入現有 Excel" accumulate correctly across multiple days."""
    if existing_raw is not None and not existing_raw.empty:
        existing_clean = _strip_raw_marker_rows(existing_raw)
        all_rows = pd.concat([existing_clean, new_rows], ignore_index=True, sort=False)
    else:
        all_rows = new_rows
    all_rows = _coalesce_amount_columns(all_rows)

    if all_rows.empty:
        return all_rows

    sort_key = all_rows[RAW_DATE_COL].astype(str)
    all_rows = all_rows.assign(_sort_key=sort_key).sort_values("_sort_key", kind="stable").drop(columns="_sort_key")

    parts = []
    for date_val, group in all_rows.groupby(RAW_DATE_COL, sort=False):
        parts.append(group)
        label = f"小計－{date_val}" if str(date_val).strip() else "小計"
        parts.append(pd.DataFrame([_raw_total_row(group, label)]))
    combined = pd.concat(parts, ignore_index=True, sort=False)
    combined = pd.concat([combined, pd.DataFrame([_raw_total_row(combined, "總計")])], ignore_index=True)
    return combined


def _extract_invoice(file_bytes: bytes, mime_type: str, translate: bool):
    model = genai.GenerativeModel("gemini-2.5-flash")
    prompt = PROMPT_AI_TRANSLATE if translate else PROMPT_AI_PLAIN
    response = model.generate_content([{"mime_type": mime_type, "data": file_bytes}, prompt])

    raw = response.text.replace("```csv", "").replace("```", "").strip()
    lines = [ln for ln in raw.split("\n") if ln.strip()]
    if not lines:
        raise ValueError("no_data")

    date_str, store_str = "", ""
    if lines[0].startswith("META###"):
        meta_parts = lines[0].split("###")
        date_str = meta_parts[1].strip() if len(meta_parts) > 1 else ""
        store_str = meta_parts[2].strip() if len(meta_parts) > 2 else ""
        lines = lines[1:]
    if not lines:
        raise ValueError("no_data")

    headers = [h.strip() for h in lines[0].split("###")]
    if len(headers) < 2:
        # Model ignored the "###" column-delimiter instruction entirely (e.g. it
        # transcribed the receipt as plain text lines instead of a table) —
        # surface this as a failed file rather than silently writing a single
        # garbled column of unrelated text into the spreadsheet.
        raise ValueError("no_data")

    rows = []
    for line in lines[1:]:
        row = [c.strip() for c in line.split("###")]
        if len(row) < len(headers):
            row += [""] * (len(headers) - len(row))
        elif len(row) > len(headers):
            row = row[:len(headers)]
        rows.append(row)

    if not rows:
        raise ValueError("no_data")

    return pd.DataFrame(rows, columns=headers), date_str, store_str


@app.route("/detect-language", methods=["POST"])
def detect_language_route():
    if not GEMINI_AVAILABLE:
        return jsonify({"language": "zh"})

    f = request.files.get("file")
    if not f:
        return jsonify({"error": "No file"}), 400

    file_bytes = f.read()
    if len(file_bytes) > GEMINI_MAX_BYTES:
        return jsonify({"language": "zh"})

    mime_type = _gemini_mime(f.filename)
    if not mime_type:
        return jsonify({"language": "zh"})

    try:
        model = genai.GenerativeModel("gemini-2.5-flash")
        prompt = ("請判斷這份文件中表格內容主要使用的語言。"
                  "只回答一個代碼，不要有其他文字或標點：zh／en／ja／ko／other")
        response = model.generate_content([{"mime_type": mime_type, "data": file_bytes}, prompt])
        code = response.text.strip().lower()
        if code not in ("zh", "en", "ja", "ko", "other"):
            code = "other"
        return jsonify({"language": code})
    except Exception:
        return jsonify({"language": "zh"})


@app.route("/exchange-rate", methods=["GET"])
def exchange_rate_route():
    from_cur = request.args.get("from", "KRW").upper()
    to_cur = request.args.get("to", "TWD").upper()
    try:
        resp = requests.get(f"https://open.er-api.com/v6/latest/{from_cur}", timeout=6)
        resp.raise_for_status()
        payload = resp.json()
        rate = payload["rates"][to_cur]
        as_of = payload.get("time_last_update_utc", "")
        return jsonify({"rate": rate, "as_of": as_of, "from": from_cur, "to": to_cur})
    except Exception as e:
        return jsonify({"error": str(e)[:200]}), 502


@app.route("/convert-ai", methods=["POST"])
def convert_ai_route():
    if not GEMINI_AVAILABLE:
        return jsonify({"error": "AI 模式未啟用（缺少 GEMINI_API_KEY）"}), 503

    files = request.files.getlist("file")
    if not files:
        return jsonify({"error": "No file"}), 400

    translate = request.form.get("translate") == "true"
    convert_currency = request.form.get("convert_currency") == "true"
    currency_from = request.form.get("currency_from", "").upper()
    exchange_rate = None
    if convert_currency:
        try:
            exchange_rate = float(request.form.get("exchange_rate", "0"))
        except ValueError:
            exchange_rate = None
        if not exchange_rate:
            convert_currency = False

    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")

    # ── Optional: merge into an existing Excel file ──
    existing_sheets = {}
    existing_label = None
    existing_file = request.files.get("existing_excel")
    if existing_file and existing_file.filename:
        try:
            existing_sheets = read_excel_sheets(existing_file.read())
            existing_label = existing_file.filename.rsplit(".", 1)[0]
        except Exception:
            existing_sheets = {}
            existing_label = None

    jobs = []
    for f in files:
        file_bytes = f.read()
        if len(file_bytes) > GEMINI_MAX_BYTES:
            jobs.append((f.filename, None, "too_large"))
            continue
        mime_type = _gemini_mime(f.filename)
        if not mime_type:
            jobs.append((f.filename, None, "bad_format"))
            continue
        jobs.append((f.filename, file_bytes, mime_type))

    had_existing_file = bool(existing_file and existing_file.filename)

    def _run(job):
        filename, file_bytes, mime_type = job
        if file_bytes is None:
            return filename, None, "", "", mime_type  # reason stored in mime_type slot
        try:
            df, date_str, store_str = _extract_invoice(file_bytes, mime_type, translate)
            return filename, df, date_str, store_str, None
        except Exception as e:
            return filename, None, "", "", ("429" if "429" in str(e) else "no_data")

    results, meta, failures = {}, {}, []
    with concurrent.futures.ThreadPoolExecutor(max_workers=min(4, len(jobs))) as ex:
        for filename, df, date_str, store_str, reason in ex.map(_run, jobs):
            if df is not None:
                results[filename] = df
                meta[filename] = (date_str, store_str)
            else:
                failures.append((filename, reason))

    if not results:
        return jsonify({"error": "no_data"}), 422

    if convert_currency:
        results = {fn: _add_ntd_column(df, exchange_rate) for fn, df in results.items()}

    used_sheet_names = set(existing_sheets.keys())
    sheets = {}
    raw_parts = []
    for filename, df in results.items():
        sheet_name = _excel_safe_sheet_name(filename.rsplit(".", 1)[0], used_sheet_names)
        sheets[sheet_name] = pd.concat(
            [df, pd.DataFrame([_per_file_subtotal_row(df)])], ignore_index=True)

        date_str, store_str = meta[filename]
        raw_parts.append(_normalize_invoice_rows(df, date_str, store_str))

    new_raw_rows = (pd.concat(raw_parts, ignore_index=True, sort=False)
                     if raw_parts else pd.DataFrame(columns=RAW_COLUMNS))

    rate_note = None
    if convert_currency:
        today = datetime.datetime.now().strftime("%Y-%m-%d")
        rate_note = f"匯率：1 {currency_from or '外幣'} = {exchange_rate} TWD（查詢日期：{today}）"

    existing_raw = existing_sheets.pop(RAW_SHEET_NAME, None) if existing_sheets else None
    raw_sheet = _build_raw_sheet(new_raw_rows, existing_raw)
    all_sheets = {RAW_SHEET_NAME: raw_sheet, **existing_sheets, **sheets}

    if had_existing_file:
        filename_out = f"{existing_label}_更新_{ts}.xlsx"
    elif len(results) > 1:
        filename_out = f"批次轉換_{len(results)}筆_{ts}.xlsx"
    else:
        only_name = next(iter(results.keys())).rsplit(".", 1)[0]
        filename_out = f"{only_name}_{ts}.xlsx"

    notes = ({name: rate_note for name, df in all_sheets.items() if NTD_COL in df.columns}
             if rate_note else None)

    excel_bytes = build_excel(all_sheets, notes=notes)
    mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

    resp = send_file(BytesIO(excel_bytes), mimetype=mime,
                     as_attachment=True, download_name=filename_out)
    resp.headers["X-Success-Count"] = str(len(results))
    resp.headers["X-Total-Count"] = str(len(jobs))
    if failures:
        resp.headers["X-Convert-Warnings"] = ";".join(
            f"{quote(fn)}:{reason}" for fn, reason in failures)
    return resp


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=False)
