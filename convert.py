import streamlit as st
import pdfplumber
import pandas as pd
import json
from io import BytesIO
from PIL import Image, ImageDraw

try:
    from pdf2image import convert_from_bytes
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from pdf2docx import Converter as PDF2Docx
    PDF2DOCX_AVAILABLE = True
except ImportError:
    PDF2DOCX_AVAILABLE = False

FORMATS = ["Excel (.xlsx)", "Word (.docx)", "Markdown (.md)", "JSON (.json)"]

# ── Custom app icon — red rounded square with white paper + arrow ──
def _pdf_icon() -> Image.Image:
    size = 512
    img = Image.new("RGBA", (size, size), (0, 0, 0, 0))
    d = ImageDraw.Draw(img)
    # Red rounded square background
    d.rounded_rectangle([0, 0, size, size], radius=110, fill="#E63946")
    # White document body
    d.rounded_rectangle([100, 70, 320, 420], radius=18, fill="#FFFFFF")
    # Folded corner (dog-ear)
    d.polygon([(254, 70), (320, 70), (320, 136), (254, 136)], fill="#E63946")
    d.polygon([(254, 70), (320, 136), (254, 136)], fill="#C1121F")
    # Text lines on document
    for y in [180, 220, 260, 300, 340]:
        d.rounded_rectangle([130, y, 290, y + 22], radius=6, fill="#E63946")
    # White arrow pointing right (= convert)
    d.polygon([(360, 210), (440, 256), (360, 302)], fill="#FFFFFF")
    d.rounded_rectangle([310, 242, 390, 270], radius=8, fill="#FFFFFF")
    return img

st.set_page_config(page_title="PDF 轉換器", page_icon=_pdf_icon(), layout="wide")

# Hide Streamlit branding
st.markdown("""
<style>
#MainMenu {visibility: hidden;}
footer {visibility: hidden;}
</style>
""", unsafe_allow_html=True)

st.title("PDF 轉換器")
st.caption("上傳 PDF，自動偵測最佳格式並立即轉換。支援文字型與掃描版（OCR）。")

# ──────────────────────────────────────────────
# Analysis (cached — runs once per file)
# ──────────────────────────────────────────────
@st.cache_data(show_spinner=False)
def analyze_pdf(pdf_bytes: bytes) -> dict:
    """Sample up to 5 pages to recommend the best output format."""
    table_pages, text_pages, total_pages = 0, 0, 0
    try:
        with pdfplumber.open(BytesIO(pdf_bytes)) as pdf:
            total_pages = len(pdf.pages)
            for page in pdf.pages[:min(5, total_pages)]:
                tables = page.extract_tables() or page.extract_tables(table_settings={
                    "vertical_strategy": "text",
                    "horizontal_strategy": "text",
                })
                if tables:
                    table_pages += 1
                if (page.extract_text() or "").strip():
                    text_pages += 1
    except Exception:
        pass

    if text_pages == 0 and table_pages == 0:
        return {
            "format": "Word (.docx)",
            "badge": "🔍 掃描版 PDF",
            "reason": "未偵測到可擷取的文字，建議啟用 OCR 再轉為 Word。",
            "use_ocr": True,
        }
    elif table_pages >= text_pages * 0.6:
        return {
            "format": "Excel (.xlsx)",
            "badge": "📊 表格型 PDF",
            "reason": f"前 {min(5, total_pages)} 頁中有 {table_pages} 頁含表格，最適合匯出為 Excel。",
            "use_ocr": False,
        }
    elif table_pages > 0:
        return {
            "format": "Word (.docx)",
            "badge": "📝 混合型 PDF",
            "reason": f"文字與表格混合（{table_pages} 頁有表格），建議轉為 Word 保留排版。",
            "use_ocr": False,
        }
    else:
        return {
            "format": "Markdown (.md)",
            "badge": "📄 純文字 PDF",
            "reason": "純文字內容，Markdown 最輕量且易讀。",
            "use_ocr": False,
        }


# ──────────────────────────────────────────────
# Extraction (cached)
# ──────────────────────────────────────────────
def _unique_cols(headers: list) -> list:
    seen: dict = {}
    out = []
    for h in headers:
        key = h if h else "col"
        if key in seen:
            seen[key] += 1
            out.append(f"{key}_{seen[key]}")
        else:
            seen[key] = 0
            out.append(key)
    return out


@st.cache_data(show_spinner=False)
def extract_tables(pdf_bytes: bytes) -> dict:
    result = {}
    with pdfplumber.open(BytesIO(pdf_bytes)) as pdf:
        for i, page in enumerate(pdf.pages):
            tables = page.extract_tables(table_settings={
                "vertical_strategy": "lines",
                "horizontal_strategy": "lines",
            })
            if not tables:
                tables = page.extract_tables(table_settings={
                    "vertical_strategy": "text",
                    "horizontal_strategy": "text",
                })
            if tables:
                dfs = []
                for t in tables:
                    cleaned = [[c if c is not None else "" for c in row] for row in t]
                    if not cleaned:
                        continue
                    if len(cleaned) > 1:
                        headers = _unique_cols(cleaned[0])
                        df = pd.DataFrame(cleaned[1:], columns=headers)
                    else:
                        df = pd.DataFrame(cleaned)
                    dfs.append(df.astype(str).replace("None", ""))
                if dfs:
                    try:
                        result[f"Page {i+1}"] = pd.concat(dfs, ignore_index=True)
                    except Exception:
                        normed = [df.rename(columns=str) for df in dfs]
                        result[f"Page {i+1}"] = pd.concat(normed, ignore_index=True)
    return result


@st.cache_data(show_spinner=False)
def extract_ocr(pdf_bytes: bytes) -> dict:
    images = convert_from_bytes(pdf_bytes, dpi=300)
    result = {}
    for i, img in enumerate(images):
        try:
            text = pytesseract.image_to_string(img, lang="chi_tra+eng")
        except pytesseract.TesseractError:
            text = pytesseract.image_to_string(img)
        if text.strip():
            result[f"Page {i+1}"] = text.strip()
    return result


# ──────────────────────────────────────────────
# Build output files
# ──────────────────────────────────────────────
def build_excel(data: dict) -> bytes:
    out = BytesIO()
    with pd.ExcelWriter(out, engine="openpyxl") as writer:
        for sheet, df in data.items():
            df.to_excel(writer, sheet_name=sheet[:31], index=False)
    return out.getvalue()


def build_word(pdf_bytes: bytes, is_ocr: bool, ocr_data: dict = None) -> bytes:
    """
    Word conversion strategy:
    - Normal PDF  → pdf2docx for high-fidelity layout recreation (fonts, tables, formatting)
    - OCR result  → plain docx with extracted text (no layout info available)
    - Fallback    → plain docx from extracted table data
    """
    if is_ocr and ocr_data:
        doc = Document()
        doc.add_heading("PDF 轉換結果（OCR）", level=0)
        for page, text in ocr_data.items():
            doc.add_heading(page, level=1)
            doc.add_paragraph(text)
        out = BytesIO()
        doc.save(out)
        return out.getvalue()

    if PDF2DOCX_AVAILABLE:
        import tempfile, os
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            f.write(pdf_bytes)
            tmp_pdf = f.name
        tmp_docx = tmp_pdf.replace(".pdf", ".docx")
        try:
            cv = PDF2Docx(tmp_pdf)
            cv.convert(tmp_docx, start=0, end=None)
            cv.close()
            with open(tmp_docx, "rb") as f:
                return f.read()
        finally:
            os.unlink(tmp_pdf)
            if os.path.exists(tmp_docx):
                os.unlink(tmp_docx)

    # Fallback if pdf2docx not available
    doc = Document()
    doc.add_heading("PDF 轉換結果", level=0)
    for page, df in (ocr_data or {}).items():
        doc.add_heading(page, level=1)
        tbl = doc.add_table(rows=1 + len(df), cols=len(df.columns))
        tbl.style = "Table Grid"
        for j, col in enumerate(df.columns):
            tbl.rows[0].cells[j].text = str(col)
        for i, row in df.iterrows():
            for j, val in enumerate(row):
                tbl.rows[i + 1].cells[j].text = str(val)
        doc.add_paragraph()
    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def build_markdown(data: dict, is_ocr: bool) -> str:
    parts = ["# PDF 轉換結果\n"]
    for page, content in data.items():
        parts.append(f"\n## {page}\n")
        parts.append(content if is_ocr else content.to_markdown(index=False))
        parts.append("\n")
    return "\n".join(parts)


def build_json(data: dict, is_ocr: bool) -> str:
    payload = data if is_ocr else {p: df.to_dict(orient="records") for p, df in data.items()}
    return json.dumps(payload, ensure_ascii=False, indent=2)


@st.cache_data(show_spinner=False)
def extract_text_pages(pdf_bytes: bytes) -> dict:
    """Extract plain text per page for editable Word output."""
    result = {}
    with pdfplumber.open(BytesIO(pdf_bytes)) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            if text.strip():
                result[f"Page {i+1}"] = text.strip()
    return result


def build_word_editable(pdf_bytes: bytes) -> bytes:
    """Clean editable Word doc — plain paragraphs, no floating boxes."""
    pages = extract_text_pages(pdf_bytes)
    doc = Document()
    for page_label, text in pages.items():
        doc.add_heading(page_label, level=1)
        for line in text.splitlines():
            if line.strip():
                doc.add_paragraph(line.strip())
        doc.add_paragraph()
    out = BytesIO()
    doc.save(out)
    return out.getvalue()


# ──────────────────────────────────────────────
# Main UI
# ──────────────────────────────────────────────
uploaded = st.file_uploader("📂 上傳 PDF", type=["pdf"])

if not uploaded:
    st.divider()
    st.markdown("上傳後系統自動偵測類型，推薦最適格式：")
    st.markdown(
        "📊 **表格型** → Excel &nbsp;&nbsp;"
        "📝 **混合型** → Word &nbsp;&nbsp;"
        "📄 **純文字** → Markdown &nbsp;&nbsp;"
        "🔍 **掃描版** → OCR + Word"
    )
    st.stop()

# ── File uploaded ─────────────────────────────
pdf_bytes = uploaded.read()
base_name = uploaded.name.rsplit(".", 1)[0]
timestamp = pd.Timestamp.now().strftime("%Y%m%d_%H%M%S")

with st.spinner("分析中..."):
    rec = analyze_pdf(pdf_bytes)

# Badge — large and prominent
st.markdown(f"## {rec['badge']}")
st.caption(rec['reason'])

st.divider()

# Format + OCR — always visible
col_fmt, col_ocr = st.columns([4, 1])
with col_fmt:
    output_format = st.radio("輸出格式", FORMATS,
                             index=FORMATS.index(rec["format"]), horizontal=True)
with col_ocr:
    if OCR_AVAILABLE:
        use_ocr = st.checkbox("🔍 OCR 模式", value=rec.get("use_ocr", False),
                              help="掃描版 PDF 請勾選，將每頁影像化後辨識文字。")
    else:
        use_ocr = False
        if rec.get("use_ocr"):
            st.warning("偵測到掃描版，但 OCR 未安裝。")

# Word mode toggle — only shown when Word is selected
word_editable = False
if output_format == "Word (.docx)" and not use_ocr:
    word_editable = st.toggle(
        "✏️ 可編輯優先（純文字段落）",
        value=False,
        help="關閉：保留原始版面（pdf2docx，部分內容可能為文字框）\n開啟：輸出為乾淨可編輯段落（無版面，但可直接在 Word 修改）"
    )

# ── Convert ───────────────────────────────────
st.divider()
with st.spinner("轉換中，請稍候..."):
    if use_ocr and OCR_AVAILABLE:
        data = extract_ocr(pdf_bytes)
        is_ocr = True
    else:
        data = extract_tables(pdf_bytes)
        is_ocr = False

if not data:
    msg = ("⚠️ 找不到表格結構。若為掃描版 PDF，請勾選上方 **OCR 模式** 再試。"
           if not use_ocr else "⚠️ OCR 未能辨識出文字，請確認 PDF 影像品質。")
    st.warning(msg)
    st.stop()

# ── Build output ──────────────────────────────
fmt = output_format

if fmt == "Excel (.xlsx)":
    file_bytes = build_excel(data)
    file_name = f"{base_name}_{timestamp}.xlsx"
    mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

elif fmt == "Word (.docx)":
    if word_editable:
        file_bytes = build_word_editable(pdf_bytes)
    else:
        file_bytes = build_word(pdf_bytes, is_ocr, ocr_data=data if is_ocr else None)
    file_name = f"{base_name}_{timestamp}.docx"
    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

elif fmt == "Markdown (.md)":
    md_text = build_markdown(data, is_ocr)
    file_bytes = md_text.encode("utf-8")
    file_name = f"{base_name}_{timestamp}.md"
    mime = "text/markdown"

elif fmt == "JSON (.json)":
    json_str = build_json(data, is_ocr)
    file_bytes = json_str.encode("utf-8")
    file_name = f"{base_name}_{timestamp}.json"
    mime = "application/json"

# ── Download ──────────────────────────────────
st.success(f"✅ 完成！已處理 {len(data)} 頁")
st.download_button(
    label=f"📥 下載 {fmt}",
    data=file_bytes,
    file_name=file_name,
    mime=mime,
    type="primary",
    use_container_width=True,
)
